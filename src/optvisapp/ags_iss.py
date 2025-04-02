"""
Module to read ISS orbit file, AGS observing plan file, and to update the latter to include more information
"""

import numpy as np
import pandas as pd
import gzip
import astropy.units as u
from astropy.time import Time

import sys
import argparse
import os

import re
import docx

import urllib.request
import shutil

from optvisapp import observing_geometry
from optvisapp.optvisapp_logging import get_logger

sys.dont_write_bytecode = True

# Log config
############
logger = get_logger(__name__)


def read_iss_oem_ephem(issorbitfile):
    """
    Read-in ISS OEM ephemeris as a dataframe
    :param issorbitfile: iss OEM ephem orbit file
    :type issorbitfile: str
    :return iss_oem_ephem: ISS OEM ephemeris
    :rtype: pandas.DataFrame
    """
    # Line after which ephemerides start
    key = 'COMMENT End sequence of events'
    with open(issorbitfile, 'r') as fp:
        lines = fp.readlines()
        for row in lines:
            # check if string present on a current line
            # find() method returns -1 if the value is not found,
            # if found it returns index of the first occurrence of the substring
            while row.find(key) != -1:
                index = lines.index(row)
                break

    ####################
    # Skip these rows and read the rest of the table as a pandas dataframe
    iss_oem_ephem = pd.read_csv(issorbitfile, sep=r"\s+", header=None, skiprows=index + 1,
                                names=["TIME_UTC", "ISS_X", "ISS_Y", "ISS_Z", "ISS_Vx", "ISS_Vy", "ISS_Vz"])
    iss_oem_ephem['TIME_UTC'] = pd.to_datetime(iss_oem_ephem['TIME_UTC'], format='ISO8601', utc=True)

    return iss_oem_ephem


def downloadissoemfile(outputdir='./'):
    """
    Download ISS OEM ephemeris from web
    :param outputdir: directory where to store the downloaded file
    :type outputdir: str
    :return iss_oem_ephem: ISS OEM ephemeris
    :rtype: str
    """
    issoemfile = 'https://nasa-public-data.s3.amazonaws.com/iss-coords/current/ISS_OEM/ISS.OEM_J2K_EPH.txt'
    iss_oem_ephem = "ISS.OEM_J2K_EPH.txt"
    urllib.request.urlretrieve(issoemfile, iss_oem_ephem)
    if not os.path.exists(outputdir):
        os.makedirs(outputdir)
        shutil.move(iss_oem_ephem, outputdir)

    return iss_oem_ephem


def read_target_catalog(targetcatalog):
    """
    Read-in a NICER target catalog
    :param targetcatalog: NICER visibility file
    :type targetcatalog: str
    :return df_nicer_catalog: NICER catalog
    :rtype: pandas.DataFrame
    """
    # Use the third row as column names
    column_names = pd.read_csv(targetcatalog, skiprows=2, nrows=1, header=None).values[0]

    # Read the remaining data, skipping the first three rows
    targetcat_df = pd.read_csv(targetcatalog, skiprows=3, header=None, names=column_names, index_col=False)

    # Remove leading/trailing whitespaces from target name
    targetcat_df['Source'] = targetcat_df['Source'].str.strip()

    # Remove duplicates from target catalog dataframe according to 'ID' column (Target ID)
    targetcat_df_nosourceduplicates = targetcat_df.drop_duplicates(subset='ID', keep='first')

    # Header of dataframe
    targetcat_header = pd.read_csv(targetcatalog, nrows=2, header=None, names=column_names, index_col=0)

    return targetcat_df, targetcat_df_nosourceduplicates, targetcat_header


def read_ags3_vis_file(ags3_vis_file):
    """
    Read-in AGS NICER visibility file as a dataframe
    :param ags3_vis_file: NICER visibility file
    :type ags3_vis_file: str
    :return df_nicer_vis: NICER visibility
    :rtype: pandas.DataFrame
    """

    key = 'Target'
    with gzip.open(ags3_vis_file, 'rt') as fp:
        lines = fp.readlines()
        for row in lines:
            # check if string present on a current line
            # find() method returns -1 if the value is not found,
            # if found it returns index of the first occurrence of the substring
            while row.find(key) != -1:
                index = lines.index(row)
                break

    ####################
    # Skip these rows and read the rest of the table as a pandas dataframe
    df_nicer_vis = pd.read_csv(ags3_vis_file, sep=r"\s+", header=None, skiprows=index + 3,
                               names=["target_name", "target_id", "vis_start", "vis_end", "Span", "Initial",
                                      "Final", "Relative"])

    df_nicer_vis['vis_start'] = pd.to_datetime(df_nicer_vis['vis_start'], format='%Y-%jT%H:%M:%S', utc=True)
    df_nicer_vis['vis_end'] = pd.to_datetime(df_nicer_vis['vis_end'], format='%Y-%jT%H:%M:%S', utc=True)

    # Remove leading/trailing whitespaces from target name
    df_nicer_vis['target_name'] = df_nicer_vis['target_name'].str.strip()

    # Drop duplicates of exact target_id and start or end of visibility windows, keep first
    mask = (df_nicer_vis.duplicated(subset=['target_id', 'vis_start']) |
            df_nicer_vis.duplicated(subset=['target_id', 'vis_end']))
    df_nicer_vis_nosrcdulpicate = df_nicer_vis[~mask]

    return df_nicer_vis, df_nicer_vis_nosrcdulpicate


def ags_update_persource(iss_oem_ephem, df_nicer_vis, targetid, srcRA, srcDEC, daysafter=1):
    """
    Calculates orbitday files
    :param iss_oem_ephem: ISS OEM ephemeris
    :type iss_oem_ephem: pandas.DataFrame
    :param df_nicer_vis: NICER visibility
    :type df_nicer_vis: pandas.DataFrame
    :param targetid: unique identifier of a target
    :type targetid: str
    :param srcRA: Right ascension in degrees J2000
    :type srcRA: float
    :param srcDEC: Declination in degrees J2000
    :type srcDEC: float
    :param daysafter: how many days after AGS start to include
    :type daysafter: float
    :return od_vis: visibility windows with min and max bright earth angle for orbit_day windows
    :rtype: pandas.DataFrame
    """
    # Filter for source
    nicer_vis_windows = df_nicer_vis[df_nicer_vis['target_id'] == targetid].reset_index(drop=True)

    # Calculate orbit information (day/night/both) for each visibility window
    nicer_vis_windows_orbit = observing_geometry.viswindow_islit(nicer_vis_windows, iss_oem_ephem)

    nicer_vis_windows_orbitday = nicer_vis_windows_orbit[nicer_vis_windows_orbit['orbit'] ==
                                                         'o_d'].reset_index(drop=True)

    # Filter the above for desirved daysafter
    nicer_vis_windows_orbitday_flt = nicer_vis_windows_orbitday[nicer_vis_windows_orbitday['vis_start'].between(
        nicer_vis_windows_orbitday['vis_start'].head(1)[0],
        nicer_vis_windows_orbitday['vis_start'].head(1)[0] + pd.Timedelta(days=daysafter))]

    # Loop over each visibility window
    nicer_vis_windows_orbitday_flt_be = []
    for ii in nicer_vis_windows_orbitday_flt.index:
        # Each full visibility window row
        nicer_vis_windows_orbitday_flt_tmp = nicer_vis_windows_orbitday_flt.loc[ii]

        # Defining visibility start and end times
        if nicer_vis_windows_orbitday_flt_tmp['Span'] < 240:
            vis_start = nicer_vis_windows_orbitday_flt_tmp['vis_start'] - pd.Timedelta(seconds=120)
            vis_end = nicer_vis_windows_orbitday_flt_tmp['vis_end'] + pd.Timedelta(seconds=120)
        else:
            vis_start = nicer_vis_windows_orbitday_flt_tmp['vis_start']
            vis_end = nicer_vis_windows_orbitday_flt_tmp['vis_end']

        # ISS orbit window that matches source visibility window
        issorbitdata_vis_window = iss_oem_ephem[iss_oem_ephem['TIME_UTC'].between(pd.Timestamp(vis_start),
                                                                                  pd.Timestamp(vis_end))]

        # Indices of the ISS orbit windows
        indices_tmp = issorbitdata_vis_window.index

        # Difference between the first iss orbit window and visibility start
        iss_vis_diff = ((issorbitdata_vis_window.head(1)['TIME_UTC'] - vis_start) > pd.Timedelta(0)).to_numpy()

        # If difference above is positive (iss orbit start time is after the start of the visibility window),
        # let's add previous row from ISS orbit
        if iss_vis_diff:
            indices = indices_tmp.append(pd.Index([indices_tmp[0] - 1])).sort_values()
            issorbitdata_add = iss_oem_ephem.loc[indices[0]]
            issorbitdata_vis_window = pd.concat([issorbitdata_vis_window, issorbitdata_add.to_frame().T],
                                                ignore_index=False).sort_index()

        ########################################################################################
        # Let's extract iss information for the times of interest and place it into a dictionary
        iss_cartesian = np.empty((len(indices), 3), dtype=float)
        iss_velocity = np.empty((len(indices), 3), dtype=float)
        iss_deltatime = np.empty(len(indices), dtype=float)
        iss_times_tmp = [None] * len(indices)

        for issorbit_row in indices:
            iss_cartesian[issorbit_row - indices.min(), :] = np.array(
                [issorbitdata_vis_window.loc[issorbit_row]['ISS_X'], issorbitdata_vis_window.loc[issorbit_row]['ISS_Y'],
                 issorbitdata_vis_window.loc[issorbit_row]['ISS_Z']])
            iss_velocity[issorbit_row - indices.min(), :] = np.array(
                [issorbitdata_vis_window.loc[issorbit_row]['ISS_Vx'],
                 issorbitdata_vis_window.loc[issorbit_row]['ISS_Vy'],
                 issorbitdata_vis_window.loc[issorbit_row]['ISS_Vz']])
            iss_times_tmp[issorbit_row - indices.min()] = issorbitdata_vis_window.loc[issorbit_row][
                'TIME_UTC'].strftime('%Y-%m-%dT%H:%M:%S')

            if issorbit_row != indices.max():
                iss_deltatime[issorbit_row - indices.min()] = (
                        issorbitdata_vis_window.loc[issorbit_row + 1]['TIME_UTC'] -
                        issorbitdata_vis_window.loc[issorbit_row]['TIME_UTC']).total_seconds()
            else:
                iss_deltatime[issorbit_row - indices.min()] = (
                        vis_end - issorbitdata_vis_window.loc[issorbit_row]['TIME_UTC']).total_seconds()

        iss_times = Time(iss_times_tmp, format='isot', scale='utc')

        iss_vis_source_info = {'ISS_coord': iss_cartesian, 'ISS_veloc': iss_velocity, 'ISStimes': iss_times,
                               'ISSdeltatime': iss_deltatime}

        #########################################################
        # Calculate bright earth angle with all above information
        numberofsteps = 4  # Think of this as 240 seconds/numberofsteps(=4) = 60 seconds; 240 seconds is the jump
        # between time-stamps in ISS orbit file. Hence, we are calculating bright_earth angle every ~minute
        t_interval = np.linspace(0, iss_vis_source_info['ISSdeltatime'][0], numberofsteps, endpoint=False)
        bright_earth_angle = np.empty(len(t_interval) * len(iss_vis_source_info['ISS_coord']), dtype=float)
        time = np.array([])

        # This is not great and needs to be written more efficiently
        counter = 0
        for kk in range(len(iss_vis_source_info['ISS_coord'])):
            for index, jj in enumerate(t_interval):
                time_tmp = iss_vis_source_info['ISStimes'][kk] + jj * u.s
                iss_cartesian_extrapolated = np.array(
                    [iss_vis_source_info['ISS_coord'][kk, 0] + iss_vis_source_info['ISS_veloc'][kk, 0] * jj,
                     iss_vis_source_info['ISS_coord'][kk, 1] + iss_vis_source_info['ISS_veloc'][kk, 1] * jj,
                     iss_vis_source_info['ISS_coord'][kk, 2] + iss_vis_source_info['ISS_veloc'][kk, 2] * jj])

                bright_earth_angle[counter] = observing_geometry.bright_earth_angle(
                    iss_cartesian=iss_cartesian_extrapolated,
                    time=time_tmp, src_ra=srcRA,
                    src_dec=srcDEC)
                time = np.append(time, time_tmp)
                # Ugh
                counter += 1

        # Putting the times and corresponding bright_earth angles together
        time_beangle = pd.DataFrame(np.vstack((time, bright_earth_angle)).T, columns=['Time', 'be_angle'])
        time_beangle['Time'] = pd.to_datetime(time_beangle['Time'], format='ISO8601', utc=True)
        # Filtering the above to match visibility window (which could start in the middle of an ISS timestamp)
        time_beangle = time_beangle[time_beangle['Time'].between(pd.Timestamp(vis_start), pd.Timestamp(vis_end))]
        # Min and max of bright-earth angles for the visibility window
        be_angle_info = pd.Series([time_beangle['be_angle'].min(), time_beangle['be_angle'].max()],
                                  index=['be_angle_min', 'be_angle_max'])
        # Adding this info to visibility window
        nicer_vis_windows_orbitday_flt_tmp_be = pd.concat([nicer_vis_windows_orbitday_flt_tmp, be_angle_info])

        # Append list with visibility windows with orbit and bright_earth information
        nicer_vis_windows_orbitday_flt_be.append(nicer_vis_windows_orbitday_flt_tmp_be)

    # Create DataFrame of the above list
    nicer_vis_windows_orbitday_flt_be_df = pd.concat(nicer_vis_windows_orbitday_flt_be, axis=1).T

    return nicer_vis_windows_orbitday_flt_be_df


# Determine whether the file is gzipped (by extension or magic number)
def is_gzipped(filepath):
    with open(filepath, 'rb') as f:
        return f.read(2) == b'\x1f\x8b'


def read_planningdoc(planningdoc):
    # Load the document
    doc = docx.Document(planningdoc)

    # List to hold processed column 2 data
    target_id_list = []

    # Loop over all tables in the document
    for table_index, table in enumerate(doc.tables):
        logger.info(f"Processing table {table_index + 1} in planning doc...")
        for row in table.rows:
            # Get text from the second cell (index 1 - traget_id)
            col2_text = get_non_strike_text(row.cells[1])

            # Skip the row if cell is empty after filtering
            if not col2_text:
                continue

            # Process col2_text: split at " or " or "\n"
            col2_text = re.split(r" or |\n", col2_text)
            col2_text = list(filter(None, col2_text))
            target_id_list.extend(col2_text)

    df_planningdoc_tragetids = pd.DataFrame(target_id_list, columns=['target_id'])
    df_planningdoc_tragetids = df_planningdoc_tragetids[~df_planningdoc_tragetids['target_id'].isin(['TargID',
                                                                                                     'OLDI',
                                                                                                     'Targ ID'])]

    df_planningdoc_tragetids['target_id'].astype(int)

    return df_planningdoc_tragetids


def get_non_strike_text(cell):
    """
    Extracts text from a .docx table cell by concatenating text runs that are not struck through
    """
    text_parts = []
    for para in cell.paragraphs:
        for run in para.runs:
            # Skip text that is struck through
            if run.font and run.font.strike:
                continue
            text_parts.append(run.text)
    return "".join(text_parts).strip()


def main():
    parser = argparse.ArgumentParser(description="Append AGS3 with bright-earth angle for orbit-day visibility windows")
    parser.add_argument("issorbitfile", help="A ISS orbit file (OEM format)", type=str)
    parser.add_argument("ags3_vis_file", help="A NICER AGS3 visibility file", type=str)
    parser.add_argument("srcname", help="Name of source for which to provide more information", type=str)
    parser.add_argument("srcRA", help="Source RA in degrees (J2000)", type=float)
    parser.add_argument("srcDEC", help="Source DEC in degrees (J2000)", type=float)
    parser.add_argument("-da", "--daysafter", help="Filtering for number of days after start of "
                                                   "visibility window (to speed things up mainly, default = 1)",
                        type=float, default=1)
    args = parser.parse_args()

    iss_oem_ephem = read_iss_oem_ephem(args.issorbitfile)
    df_nicer_vis = read_ags3_vis_file(args.ags3_vis_file)

    ags_update_persource(iss_oem_ephem, df_nicer_vis, args.srcname, args.srcRA, args.srcDEC, args.daysafter)

    return


if __name__ == '__main__':
    main()
